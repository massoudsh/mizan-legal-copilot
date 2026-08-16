"""تست‌های استخراج متن از فرمت‌های مختلف فایل (document_parser)."""

import io
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402
from pypdf import PdfWriter  # noqa: E402

from app.document_parser import UnsupportedFileTypeError, extract_text  # noqa: E402


def test_extract_txt():
    assert extract_text("note.txt", "سلام دنیا".encode("utf-8")) == "سلام دنیا"


def test_extract_md():
    assert extract_text("note.md", b"# Title\n\nBody") == "# Title\n\nBody"


def test_extract_txt_case_insensitive_extension():
    assert extract_text("NOTE.TXT", b"hi") == "hi"


def test_extract_docx():
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("خط اول")
    doc.add_paragraph("خط دوم")
    doc.save(buf)

    text = extract_text("file.docx", buf.getvalue())

    assert "خط اول" in text
    assert "خط دوم" in text


def test_extract_pdf_blank_page_returns_empty_text():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)

    text = extract_text("file.pdf", buf.getvalue())

    assert text == ""


def test_unsupported_extension_raises():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text("file.xyz", b"content")
